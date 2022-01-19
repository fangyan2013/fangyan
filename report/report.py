import numpy as np
import pandas as pd
from matplotlib import pyplot as plt
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
from docx2pdf import convert
import json
class report:
    def __init__(self,json_data):
        self.json_data=json_data
        self.workfliow_starus_list=None
        self.workfliow_name_dict=None
    
    
    
    
    def generate_picture(self):
        '''生成图片'''
        labels = []
        men_means = []
        women_means = []
        for i in self.workfliow_starus_list:
            labels.append(list(i.keys())[0])
            men_means.append(list(i.values())[0]['error_len'])
            women_means.append(list(i.values())[0]['Success_len'])
            
        print(labels,men_means,women_means)

        width = 0.35       # the width of the bars: can also be len(x) sequence

        fig, ax = plt.subplots()

        ax.bar(labels, men_means, width, label='success')
        ax.bar(labels, women_means, width, bottom=men_means,label='error')

        ax.set_ylabel('many')
        ax.set_title('Success and failure by workflow')
        ax.legend()
        plt.savefig('./test.jpg')


    def generate_word(self):
        '''生成word'''
        # 新建文档
        doc2 = Document()
        # 新增文档标题
        doc2.add_heading('报告',0)
        doc2.add_heading('图片',2)
        # 增加图像
        doc2.add_picture('test.jpg', width=Inches(5.5))
        doc2.add_heading('表格',2)
        # 增加表格，这是表格头
        table = doc2.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '编号'
        hdr_cells[1].text = '工作流'
        hdr_cells[2].text = '状态'
        # 这是表格数据
        
        
        
        
        # 遍历数据并展示
        for id in list(range(len(self.workfliow_starus_list))):
            row_cells = table.add_row().cells
            row_cells[0].text = str(id)
            row_cells[1].text = self.workfliow_starus_list[id].keys()
            if list(self.workfliow_starus_list[id].values())[0]['error_len']>0:
                row_cells[2].text = "error"
            else:
                row_cells[2].text = "Success"



        doc2.add_heading('日志',2)
        # 增加有序列表
        
        for id in self.json_data['log']:
            self.workfliow_name_dict
        
            Workflow_name = self.workfliow_name_dict[id]
            doc2.add_paragraph(
                Workflow_name.strip(), style='List Number'
            )
            
            node_obj = self.json_data['log'][id]
            if isinstance(node_obj,str):
                node_obj=json.loads(self.json_data['workfliow'][id]['Workflwow_obj'])
            for node_id in node_obj:
                doc2.add_paragraph(
                    "节点名称:%s" % node_obj[node_id]['Workflow_name'].strip(), style='List Bullet'
                )
                doc2.add_paragraph(
                    "状态:%s" % node_obj[node_id]['status']
                ).paragraph_format.left_indent = Pt(20)
                doc2.add_paragraph(
                    "日志:%s" % node_obj[node_id]['Return_parameter']
                ).paragraph_format.left_indent = Pt(20)
            
            
        # doc2.add_paragraph(
            # '工作流', style='List Number'
        # )
        # doc2.add_paragraph(
            # '节点二', style='List Bullet'
        # )
        # doc2.add_paragraph(
            # '节点三', style='List Bullet'
        # )
        # doc2.add_paragraph(
            # '节点四', style='List Bullet'
        # )

        # doc2.add_paragraph(
            # '工作流二', style='List Number'
        # )

        # doc2.add_paragraph(
            # '苹果', style='List Bullet'
        # )
        # doc2.add_paragraph(
            # '香蕉', style='List Bullet'
        # )
        # doc2.add_paragraph(
            # '馄炖', style='List Bullet'
        # )
        # doc2.add_paragraph(
            # '减少加班时间', style='List Number'
        # )
        doc2.save('word2.docx')


    def Conversion_pdf(self):
        '''转换pdf'''
        convert("./word2.docx")

    def data_processing(self):
        '''参数处理'''
        workfliow_name_dict={}
        for id in self.json_data['workfliow']:
            workfliow_name_dict[id]=self.json_data['workfliow'][id]['Workflow_name']
        self.workfliow_name_dict=workfliow_name_dict
        
        print(workfliow_name_dict,'workfliow_name_dict')
        
        workfliow_id_list= []
        workfliow_id_list_error =[]
        for workfliow_id in self.json_data['log']:
            workfliow_id_list.append(workfliow_name_dict[workfliow_id])
            for node_id in self.json_data['log'][workfliow_id]:
                if self.json_data['log'][workfliow_id][node_id]['status']=="error":
                    #失败的节点
                    workfliow_id_list_error.append(workfliow_name_dict[workfliow_id])
                    break
        workfliow_id_list_Success= list(set(workfliow_id_list).difference(set(workfliow_id_list_error)))
        
        print(workfliow_id_list,workfliow_id_list_error,workfliow_id_list_Success,'你问我')
        #--------------------------------------------------------------------------------
        self.workfliow_id_list=workfliow_id_list
        self.workfliow_id_list_error=workfliow_id_list_error
        self.workfliow_id_list_Success=workfliow_id_list_Success
        
        
        workfliow_starus_list=[]
        for workfliow_id in self.json_data['log']:
            workfliow_dict={}
            error_len=0;Success_len=0
            for node_id in self.json_data['log'][workfliow_id]:
                if self.json_data['log'][workfliow_id][node_id]['status']=="error":
                    #失败的节点
                    error_len=error_len+1
                else:
                    Success_len=Success_len+1
            #workfliow_dict[workfliow_id]={
            workfliow_dict[workfliow_name_dict[workfliow_id]]={
                "error_len":error_len,
                "Success_len":Success_len
            }
            workfliow_starus_list.append(workfliow_dict)
        print(workfliow_starus_list,'图标数据')
        self.workfliow_starus_list=workfliow_starus_list
        # self.generate_picture()#生成图表
        # self.generate_word()#生成word
        # self.Conversion_pdf()#转pdf
        
        #--------------------------------------------------------------------------------

    def run_report(self):
        '''执行'''
        self.data_processing()
        return {"workfliow_data":self.json_data,"workfliow":self.workfliow_starus_list,"chart":self.workfliow_name_dict}